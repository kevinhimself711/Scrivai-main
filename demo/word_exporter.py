"""将生成的 Markdown 转换为 Word (.docx) 字节流。

专为 Scrivai 施工方案场景设计：处理标题、段落、Markdown 表格。
足以覆盖当前模板输出的 99% 内容。
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
_TABLE_ROW_PATTERN = re.compile(r"^\|(.+)\|\s*$")
_TABLE_SEPARATOR_PATTERN = re.compile(r"^\|[\s\-:|]+\|\s*$")


def markdown_to_docx(markdown_text: str, *, title: str | None = None) -> bytes:
    """把 Markdown 文本转换成 DOCX 二进制。

    支持：`# ~ ######` 标题、Markdown 管道表格、普通段落。
    """
    document = Document()

    _configure_default_style(document)

    if title:
        heading = document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown_text.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx]

        if not line.strip():
            idx += 1
            continue

        heading_match = _HEADING_PATTERN.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            document.add_heading(text, level=min(level, 9))
            idx += 1
            continue

        # 尝试识别表格（至少需要表头行 + 分隔行）
        if _TABLE_ROW_PATTERN.match(line.strip()):
            table_lines: list[str] = []
            lookahead = idx
            while lookahead < len(lines) and _TABLE_ROW_PATTERN.match(
                lines[lookahead].strip()
            ):
                table_lines.append(lines[lookahead].strip())
                lookahead += 1
            if (
                len(table_lines) >= 2
                and _TABLE_SEPARATOR_PATTERN.match(table_lines[1])
            ):
                _add_markdown_table(document, table_lines)
                idx = lookahead
                continue

        document.add_paragraph(line)
        idx += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _configure_default_style(document: Document) -> None:
    """设置默认中文字体，避免 Word 里中文变方块。"""
    normal_style = document.styles["Normal"]
    font = normal_style.font
    font.name = "宋体"
    font.size = Pt(11)
    # 让中文字体在 East Asia 字符集生效
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")


def _add_markdown_table(document: Document, table_lines: list[str]) -> None:
    """解析 Markdown 管道表格并写入 DOCX。"""
    header_cells = _split_row(table_lines[0])
    data_rows: list[list[str]] = []
    for line in table_lines[2:]:  # 跳过 header 行和分隔行
        cells = _split_row(line)
        if len(cells) < len(header_cells):
            cells.extend([""] * (len(header_cells) - len(cells)))
        data_rows.append(cells[: len(header_cells)])

    table = document.add_table(rows=1, cols=len(header_cells))
    table.style = "Table Grid"

    header_row = table.rows[0].cells
    for col_idx, text in enumerate(header_cells):
        header_row[col_idx].text = text

    for row_data in data_rows:
        row_cells = table.add_row().cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text


def _split_row(line: str) -> list[str]:
    """拆解一行 Markdown 表格，去掉首尾竖线。"""
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]
